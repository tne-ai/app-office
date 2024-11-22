import markdown
from markdown.extensions.tables import TableExtension
from docx import Document
from bs4 import BeautifulSoup

def markdown_to_docx(md_file, output_file):
    """
    Convert a Markdown file to a Word document.

    :param md_file: Path to the input Markdown file.
    :param output_file: Path to save the Word document.
    """
    try:
        # Read the Markdown file
        with open(md_file, 'r', encoding='utf-8') as f:
            md_content = f.read()

        # Convert Markdown to HTML
        html_content = markdown.markdown(md_content, extensions=['extra', TableExtension()])

        # Use BeautifulSoup to parse the HTML
        soup = BeautifulSoup(html_content, 'html.parser')

        # save the html
        with open('html/inventory-reco.html', 'w', encoding='utf-8') as f:
            f.write(soup.prettify())

        # Create a Word document
        doc = Document()

        # Parse HTML and add content to the Word document
        for element in soup:
            if element.name == "h1":
                doc.add_heading(element.text, level=1)
            elif element.name == "h2":
                doc.add_heading(element.text, level=2)
            elif element.name == "h3":
                doc.add_heading(element.text, level=3)
            elif element.name == "p":
                doc.add_paragraph(element.text)
            elif element.name == "ul":
                for li in element.find_all("li"):
                    doc.add_paragraph(li.text, style="List Bullet")
            elif element.name == "ol":
                for li in element.find_all("li"):
                    doc.add_paragraph(li.text, style="List Number")
            elif element.name == "table":
                rows = element.find_all("tr")
                if rows:
                    table = doc.add_table(rows=len(rows), cols=len(rows[0].find_all(["td", "th"])))

                    table.style = "Table Grid"
                    for i, row in enumerate(rows):
                        cells = row.find_all(["td", "th"])
                        for j, cell in enumerate(cells):
                            table.cell(i, j).text = cell.get_text(strip=True)

        # Save the Word document
        doc.save(output_file)
        print(f"Markdown file '{md_file}' has been converted to Word document '{output_file}'.")
    except Exception as e:
        print(f"Error: {e}")


# Example usage
md_file = "data/inventory-reco.md"  # Replace with the input Markdown file path
output_file = "docx/inventory-reco.docx"  # Replace with the output Word document path
markdown_to_docx(md_file, output_file)